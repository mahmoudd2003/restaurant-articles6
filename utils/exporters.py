from docx import Document
from docx.shared import Pt
import json
from datetime import datetime
from io import BytesIO

def to_docx(markdown_text: str) -> bytes:
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)
    for line in markdown_text.splitlines():
        if line.startswith("# "):
            p = doc.add_paragraph()
            run = p.add_run(line.replace("# ", ""))
            run.bold = True
            run.font.size = Pt(16)
        elif line.startswith("## "):
            p = doc.add_paragraph()
            run = p.add_run(line.replace("## ", ""))
            run.bold = True
            run.font.size = Pt(14)
        elif line.startswith("### "):
            p = doc.add_paragraph()
            run = p.add_run(line.replace("### ", ""))
            run.bold = True
            run.font.size = Pt(12)
        else:
            doc.add_paragraph(line)
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.read()

def to_json(structured_obj: dict) -> str:
    structured_obj["exported_at"] = datetime.utcnow().isoformat() + "Z"
    return json.dumps(structured_obj, ensure_ascii=False, indent=2)
